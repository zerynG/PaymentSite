from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.conf import settings
from .models import CommercialProposal
from .forms import CommercialProposalForm, ServiceItemFormSet
from openpyxl import Workbook
from docx import Document
from datetime import datetime
from io import BytesIO
from xhtml2pdf import pisa
from django.urls import reverse

def debug_urls(request):
    """Временная функция для отладки URL"""
    urls_to_test = [
        'commercial_proposal:proposal_list',
        'commercial_proposal:create_proposal',
        'commercial_proposal:proposal_detail',
    ]

    result = "<h1>Debug URLs</h1>"
    for url_name in urls_to_test:
        try:
            url = reverse(url_name)
            result += f"<p>✓ {url_name} -> {url}</p>"
        except Exception as e:
            result += f"<p>✗ {url_name} -> ERROR: {e}</p>"

    return HttpResponse(result)


def proposal_list(request):
    """Список всех коммерческих предложений"""
    proposals = CommercialProposal.objects.all().order_by('-creation_date')
    return render(request, 'commercial_proposal/list.html', {'proposals': proposals})

def create_proposal(request):
    if request.method == 'POST':
        form = CommercialProposalForm(request.POST)
        formset = ServiceItemFormSet(request.POST)

        if form.is_valid() and formset.is_valid():
            proposal = form.save()
            formset.instance = proposal
            formset.save()
            return redirect('commercial_proposal:proposal_detail', pk=proposal.pk)  # ИСПРАВЛЕНО
    else:
        form = CommercialProposalForm()
        formset = ServiceItemFormSet()

    return render(request, 'commercial_proposal/create.html', {
        'form': form,
        'formset': formset
    })


def proposal_detail(request, pk):
    proposal = get_object_or_404(CommercialProposal, pk=pk)
    return render(request, 'commercial_proposal/detail.html', {'proposal': proposal})


def delete_proposal(request, pk):
    """Удаление коммерческого предложения с подтверждением."""
    proposal = get_object_or_404(CommercialProposal, pk=pk)
    if request.method == "POST":
        proposal.delete()
        return redirect('commercial_proposal:proposal_list')
    return render(request, 'commercial_proposal/confirm_delete.html', {'proposal': proposal})


def download_pdf(request, pk):
    proposal = get_object_or_404(CommercialProposal, pk=pk)
    html = render_to_string('commercial_proposal/template_pdf.html', {'proposal': proposal})

    # Создаем PDF
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)

    if not pdf.err:
        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="proposal_{pk}.pdf"'
        return response

    return HttpResponse("Ошибка при создании PDF", status=500)


def download_excel(request, pk):
    proposal = get_object_or_404(CommercialProposal, pk=pk)

    wb = Workbook()
    ws = wb.active
    ws.title = "Коммерческое предложение"

    # Заголовок
    ws['A1'] = proposal.title
    ws['A2'] = f"Дата формирования: {proposal.creation_date}"
    ws['A3'] = f"Заказчик: {proposal.customer.name}"

    # Услуги
    ws['A5'] = "Услуги"
    ws['A6'] = "Название"
    ws['B6'] = "Часы"
    ws['C6'] = "Период"
    ws['D6'] = "Стоимость"

    for i, service in enumerate(proposal.services.all(), 7):
        ws[f'A{i}'] = service.name
        ws[f'B{i}'] = float(service.hours)
        ws[f'C{i}'] = f"{service.start_date} - {service.end_date}"
        ws[f'D{i}'] = float(service.cost)

    # Итог
    last_row = 7 + len(proposal.services.all())
    ws[f'A{last_row}'] = "ИТОГО:"
    ws[f'D{last_row}'] = float(proposal.total_cost)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="proposal_{pk}.xlsx"'
    wb.save(response)
    return response


def download_word(request, pk):
    proposal = get_object_or_404(CommercialProposal, pk=pk)

    doc = Document()
    doc.add_heading(proposal.title, 0)
    doc.add_paragraph(f"Дата формирования: {proposal.creation_date}")
    doc.add_paragraph(f"Заказчик: {proposal.customer.name}")

    doc.add_heading('Техническое задание', level=1)
    doc.add_paragraph(proposal.technical_spec)

    doc.add_heading('Услуги', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название'
    hdr_cells[1].text = 'Часы'
    hdr_cells[2].text = 'Период'
    hdr_cells[3].text = 'Стоимость'

    for service in proposal.services.all():
        row_cells = table.add_row().cells
        row_cells[0].text = service.name
        row_cells[1].text = str(service.hours)
        row_cells[2].text = f"{service.start_date} - {service.end_date}"
        row_cells[3].text = str(service.cost)

    doc.add_paragraph(f"ИТОГО: {proposal.total_cost} руб.")
    doc.add_paragraph(f"{proposal.manager_position} {proposal.manager_name}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="proposal_{pk}.docx"'
    doc.save(response)
    return response